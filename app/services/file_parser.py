"""File parsing service for TXT, PDF, DOCX, MD."""
from pathlib import Path
from io import BytesIO
import re
import mistune


class FileParser:
    """Unified file parsing interface."""

    SUPPORTED = {".txt", ".pdf", ".docx", ".md"}

    @classmethod
    def parse(cls, file_path: Path) -> str:
        """Parse file and return text content."""
        suffix = file_path.suffix.lower()

        if suffix == ".txt":
            return cls._parse_txt(file_path)
        elif suffix == ".pdf":
            return cls._parse_pdf(file_path)
        elif suffix == ".docx":
            return cls._parse_docx(file_path)
        elif suffix == ".md":
            return cls._parse_md(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @classmethod
    def parse_bytes(cls, content: bytes, filename: str) -> str:
        """Parse file from bytes (for uploads)."""
        suffix = Path(filename).suffix.lower()

        if suffix == ".txt":
            return content.decode("utf-8", errors="replace")
        elif suffix == ".pdf":
            return cls._parse_pdf_bytes(content)
        elif suffix == ".docx":
            return cls._parse_docx_bytes(content)
        elif suffix == ".md":
            return cls._parse_md_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _parse_txt(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        import PyPDF2
        with path.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(
                page.extract_text() or "" for page in reader.pages
            )

    @staticmethod
    def _parse_pdf_bytes(content: bytes) -> str:
        import PyPDF2
        reader = PyPDF2.PdfReader(BytesIO(content))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )

    @staticmethod
    def _parse_docx(path: Path) -> str:
        from docx import Document
        doc = Document(path)
        return "\n".join(para.text for para in doc.paragraphs)

    @staticmethod
    def _parse_docx_bytes(content: bytes) -> str:
        from docx import Document
        doc = Document(BytesIO(content))
        return "\n".join(para.text for para in doc.paragraphs)

    @staticmethod
    def _parse_md(path: Path) -> str:
        """Parse markdown file to plain text."""
        text = path.read_text(encoding="utf-8", errors="replace")
        return FileParser._md_to_text(text)

    @staticmethod
    def _parse_md_bytes(content: bytes) -> str:
        """Parse markdown bytes to plain text."""
        text = content.decode("utf-8", errors="replace")
        return FileParser._md_to_text(text)

    @staticmethod
    def _md_to_text(text: str) -> str:
        """Convert markdown to plain text using mistune."""
        # Use mistune to parse and strip markdown formatting
        markdown = mistune.create_markdown()
        html = markdown(text)
        # Simple HTML to text conversion
        text = re.sub(r'<[^>]+>', '', html)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
